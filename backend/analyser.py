import anthropic
import json
import re


client = anthropic.Anthropic()
MODEL = "claude-sonnet-4-6"


def clean_json(raw: str) -> str:
    return re.sub(r"```json|```", "", raw).strip()


def scrape_skills(cv_text: str) -> dict:
    prompt = f"""From the following CV, extract a comprehensive categorised list of skills and capabilities.
Group them under clear headings (e.g. Product Management, Technical Tools, Domain Expertise, Soft Skills, Languages).
Return ONLY a JSON object with this structure:
{{"categories": [{{"name": "Category Name", "skills": ["skill1", "skill2"]}}]}}

CV TEXT:
{cv_text}"""

    message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = message.content[0].text
    print("DEBUG analyse_match response:", raw[:500])
    return json.loads(clean_json(raw))


def analyse_match(cv_text: str, jd_text: str) -> dict:
    prompt = f"""You are an expert recruiter and senior PM career coach.
Analyse this CV against the job description and return ONLY a JSON object with this exact structure:
{{
  "overall_score": <0-100>,
  "keyword_score": <0-100>,
  "experience_score": <0-100>,
  "matches": [{{"item": "description", "evidence": "where in CV", "importance": "high|medium|low"}}],
  "partials": [{{"item": "description", "gap": "what is missing", "importance": "high|medium|low"}}],
  "gaps": [{{"item": "description", "reason": "why it is a gap", "importance": "high|medium|low"}}],
  "summary": "2-3 sentence overall assessment"
}}
Limit matches, partials, and gaps to 5 items each maximum. Keep all strings under 100 characters.
Return ONLY the JSON, no markdown, no preamble.

JOB DESCRIPTION:
{jd_text}

CV:
{cv_text}"""

    message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = message.content[0].text
    return json.loads(clean_json(raw))


def reorder_cv(cv_text: str, jd_text: str) -> dict:
    base_prompt = f"""You are an expert CV writer. Rewrite the CV below with experience bullets reordered within each role so the most relevant to the job description appear first. Do not add, invent, or remove any experience — only reorder bullets within each existing role. Keep all other sections unchanged. Keep bullet points to 20 words maximum.

JOB DESCRIPTION:
{jd_text}

CV:
{cv_text}"""

    # Call 1 – Markdown
    md_message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": base_prompt + "\n\nReturn ONLY the reordered CV as clean markdown. No preamble, no explanation."}]
    )
    markdown = md_message.content[0].text.strip()

    # Call 2 – HTML
    html_message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": base_prompt + "\n\nReturn ONLY a complete, self-contained, print-ready HTML CV with clean professional styling, navy and white colour scheme, no external dependencies. No preamble, no explanation, just the HTML."}]
    )
    html = html_message.content[0].text.strip()

    return {"html": html, "markdown": markdown}


def generate_docx(cv_text: str, jd_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    reordered = reorder_cv(cv_text, jd_text)
    md = reordered.get("markdown", cv_text)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.875)
        section.right_margin = Inches(0.875)

    NAVY = RGBColor(0x1B, 0x3A, 0x6B)
    DARK = RGBColor(0x1A, 0x1A, 0x1A)
    MID  = RGBColor(0x55, 0x55, 0x55)

    lines = md.strip().split("\n")
    for line in lines:
        line = line.rstrip()

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip().upper())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = NAVY

        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = NAVY
            pf = p.paragraph_format
            pf.space_before = Pt(10)
            pf.space_after = Pt(4)
            p.paragraph_format.border_bottom_color = NAVY

        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK

        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK

        elif line.strip() == "" or line.strip() == "---":
            doc.add_paragraph()

        else:
            p = doc.add_paragraph()
            run = p.add_run(line.strip())
            run.font.size = Pt(9.5)
            run.font.color.rgb = MID

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
