import anthropic
import json
import re


client = anthropic.Anthropic()
MODEL = "claude-sonnet-4-6"


def clean_json(raw: str) -> str:
    return re.sub(r"```json|```", "", raw).strip()


def scrape_skills(cv_text: str) -> dict:
    prompt = f"""From the following CV, extract a comprehensive categorised list of skills and capabilities.
Group them under clear headings (e.g. Product Management, Technical Tools, Domain Expertise, Soft Skills, Languages).
Return ONLY a JSON object with this structure:
{{"categories": [{{"name": "Category Name", "skills": ["skill1", "skill2"]}}]}}

CV TEXT:
{cv_text}"""

    message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = message.content[0].text
    return json.loads(clean_json(raw))


def analyse_match(cv_text: str, jd_text: str) -> dict:
    prompt = f"""You are an expert recruiter and senior PM career coach.
Analyse this CV against the job description and return ONLY a JSON object with this exact structure:
{{
  "overall_score": <0-100>,
  "keyword_score": <0-100>,
  "experience_score": <0-100>,
  "matches": [{{"item": "description", "evidence": "where in CV", "importance": "high|medium|low"}}],
  "partials": [{{"item": "description", "gap": "what is missing", "importance": "high|medium|low"}}],
  "gaps": [{{"item": "description", "reason": "why it is a gap", "importance": "high|medium|low"}}],
  "summary": "2-3 sentence overall assessment"
}}
Return ONLY the JSON, no markdown, no preamble.

JOB DESCRIPTION:
{jd_text}

CV:
{cv_text}"""

    message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = message.content[0].text
    return json.loads(clean_json(raw))


def reorder_cv(cv_text: str, jd_text: str) -> dict:
    base_prompt = f"""You are an expert CV writer. Rewrite the CV below with experience bullets reordered within each role so the most relevant to the job description appear first. Do not add, invent, or remove any experience — only reorder bullets within each existing role. Keep all other sections unchanged. Keep bullet points to 20 words maximum.

STRICT MARKDOWN FORMATTING RULES — follow exactly:
- Name: # FIRSTNAME LASTNAME (all caps)
- Contact line: plain text with | separators e.g. phone | email | linkedin | github
- Section headers: ## SECTION NAME (all caps)
- Job titles: ### Job Title  |  Company Name[TAB]Date Range
- Sub-label under job title: *Brief platform description* (single asterisks, max 8 words)
- Experience bullets: - bullet text (standard markdown bullets)
- Areas of Expertise: ONE plain text line with | separators — NO bullets NO dashes
- Skills: **Label:**[TAB]Value (one per line, bold label with colon, tab, plain value)
- Education: plain text one entry per line as: Degree  ·  Institution — NO bullets NO dashes
- No blank lines between job title, sub-label and bullets

JOB DESCRIPTION:
{jd_text}

CV:
{cv_text}"""

    md_message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": base_prompt + "\n\nReturn ONLY the reordered CV as clean markdown following the formatting rules above. No preamble, no explanation."}]
    )
    markdown = md_message.content[0].text.strip()

    html_message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": base_prompt + "\n\nReturn ONLY a complete, self-contained, print-ready HTML CV with clean professional styling, navy and white colour scheme, no external dependencies. No preamble, no explanation, just the HTML."}]
    )
    html = html_message.content[0].text.strip()

    return {"html": html, "markdown": markdown}


def generate_docx(cv_text: str, jd_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Mm
    from io import BytesIO
    import re as _re

    reordered = reorder_cv(cv_text, jd_text)
    md = reordered.get("markdown", cv_text)

    doc = Document()

    # Ensure Hyperlink style exists with correct blue/underline formatting
    from docx.oxml.ns import nsmap as _nsmap
    styles = doc.styles
    try:
        hl_style = styles['Hyperlink']
    except KeyError:
        from docx.oxml import OxmlElement as _OXE
        from docx.oxml.ns import qn as _qn
        hl_style = styles.add_style('Hyperlink', 1)  # 1 = character style
    # Set blue colour and underline on the style
    from docx.oxml import OxmlElement as _OXE
    from docx.oxml.ns import qn as _qn
    rPr = hl_style.element.get_or_add_rPr()
    # Remove existing colour/underline if any
    for tag in ['w:color', 'w:u']:
        existing = rPr.find(f'{{{_nsmap["w"]}}}{tag.split(":")[1]}')
        if existing is not None:
            rPr.remove(existing)
    color_el = _OXE('w:color')
    color_el.set(_qn('w:val'), '0563C1')  # Standard Word hyperlink blue
    u_el = _OXE('w:u')
    u_el.set(_qn('w:val'), 'single')
    rPr.append(color_el)
    rPr.append(u_el)

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # Colours matching template
    TEAL   = RGBColor(0x1D, 0x6E, 0x72)   # section headers #1D6E72
    DARK   = RGBColor(0x00, 0x00, 0x00)   # body text black
    MID    = RGBColor(0x55, 0x55, 0x55)   # secondary text #555555

    def add_bottom_border(paragraph, color="1D6E72"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def set_font(run, size, color, bold=False, italic=False):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    # Right-align tab stop at page width minus margins
    PAGE_WIDTH_TWIPS = int((8.5 - 0.9 - 0.9) * 1440)  # in twips

    def add_right_tab(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(PAGE_WIDTH_TWIPS))
        tabs.append(tab)
        pPr.append(tabs)

    # ── Numbering definition for bullets (Symbol font &#61623;) ──────────────
    from docx.oxml.ns import nsmap
    from copy import deepcopy

    def add_bullet_numbering(doc):
        """Add a numbering definition matching the template bullet style."""
        from docx.opc.part import Part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import lxml.etree as etree

        nsuri = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        
        # Check if numbering part exists
        try:
            num_part = doc.part.numbering_part
        except Exception:
            # Create numbering part if it doesn't exist
            from docx.parts.numbering import NumberingPart
            num_part = NumberingPart.new()
            doc.part.relate_to(num_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')

        tree = num_part._element

        # Add abstract numbering definition
        abstract_xml = f"""<w:abstractNum xmlns:w="{nsuri}" w:abstractNumId="10">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="bullet"/>
    <w:lvlText w:val="&#61623;"/>
    <w:lvlJc w:val="start"/>
    <w:pPr>
      <w:tabs><w:tab w:val="num" w:pos="0"/></w:tabs>
      <w:ind w:start="540" w:hanging="360"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
    </w:rPr>
  </w:lvl>
</w:abstractNum>"""

        num_xml = f"""<w:num xmlns:w="{nsuri}" w:numId="10">
  <w:abstractNumId w:val="10"/>
</w:num>"""

        abstract_elem = etree.fromstring(abstract_xml)
        num_elem = etree.fromstring(num_xml)

        # Insert before first num element or append
        nums = tree.findall(f'{{{nsuri}}}num')
        if nums:
            tree.insert(list(tree).index(nums[0]), abstract_elem)
        else:
            tree.append(abstract_elem)
        tree.append(num_elem)

        return 10  # return numId

    def apply_bullet(paragraph, num_id=10):
        """Apply bullet numbering to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), str(num_id))
        numPr.append(ilvl)
        numPr.append(numId_elem)
        pPr.append(numPr)

    bullet_num_id = add_bullet_numbering(doc)

    # ── Section tracking ────────────────────────────────────────────────────
    in_education = False
    in_skills = False
    in_expertise = False

    lines = md.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # H1 – Name (centred, black, 18pt bold, SA=3)
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[2:].strip().upper())
            set_font(run, 18, DARK, bold=True)
            in_education = False
            in_skills = False

        # Contact line (centred, 9pt MID, SA=10) with hyperlinks for email/linkedin/github
        elif "|" in line and not line.startswith("#") and i < 6:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(10)

            def add_hyperlink(paragraph, text, url):
                """Add a clickable hyperlink run to a paragraph."""
                part = paragraph.part
                r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", r_id)
                new_run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                # Apply all formatting directly – don't rely on Hyperlink style
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "Calibri")
                rFonts.set(qn("w:hAnsi"), "Calibri")
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "0563C1")
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "18")
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "18")
                rPr.extend([rFonts, color, u, sz, szCs])
                new_run.append(rPr)
                t = OxmlElement("w:t")
                t.text = text
                new_run.append(t)
                hyperlink.append(new_run)
                paragraph._p.append(hyperlink)

            parts = _re.split(r"(\s*\|\s*)", line)
            for part in parts:
                part_clean = part.strip()
                # Detect and hyperlink email, linkedin, github
                if "@" in part_clean and "." in part_clean:
                    add_hyperlink(p, part_clean, "mailto:" + part_clean)
                elif "linkedin.com" in part_clean.lower():
                    url = part_clean if part_clean.startswith("http") else "https://" + part_clean
                    add_hyperlink(p, part_clean, url)
                elif "github.com" in part_clean.lower():
                    url = part_clean if part_clean.startswith("http") else "https://" + part_clean
                    add_hyperlink(p, part_clean, url)
                else:
                    run = p.add_run(part)
                    set_font(run, 9, MID)

        # H2 – Section headers (teal 11pt bold, SB=12, SA=3, bottom border)
        elif line.startswith("## "):
            section_name = line[3:].strip().upper()
            in_education = "EDUCATION" in section_name
            in_skills = "SKILL" in section_name
            in_expertise = "EXPERTISE" in section_name
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(section_name)
            set_font(run, 11, TEAL, bold=True)
            add_bottom_border(p)

        # H3 – Job title "Bold Title  |  MID Company\tMID Date" (SB=8, SA=2)
        elif line.startswith("### "):
            in_education = False
            in_skills = False
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            add_right_tab(p)

            if "\t" in text:
                main, date = text.split("\t", 1)
            else:
                main, date = text, ""

            if "  |  " in main:
                title, company = main.split("  |  ", 1)
            elif " | " in main:
                title, company = main.split(" | ", 1)
            else:
                title, company = main, ""

            run1 = p.add_run(title.strip())
            set_font(run1, 10, DARK, bold=True)
            if company:
                run2 = p.add_run("  |  " + company.strip())
                set_font(run2, 10, MID)
            if date:
                run3 = p.add_run("\t" + date.strip())
                set_font(run3, 10, MID)

        # Sub-label – italic 9pt MID (asterisk or underscore, single not double)
        elif ((line.startswith("*") and line.endswith("*") and not line.startswith("**")) or
              (line.startswith("_") and line.endswith("_") and not line.startswith("__"))):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line.strip("*_"))
            set_font(run, 9, MID, italic=True)

        # Bullets (List Paragraph, Symbol bullet, SB=0, SA=3) – never in education or expertise
        elif (line.startswith("- ") or line.startswith("* ")) and not in_education and not in_expertise:
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            apply_bullet(p, bullet_num_id)
            run = p.add_run(line[2:].strip())
            set_font(run, 10, DARK)

        # Expertise section – collect ALL expertise lines, join with | and render as one paragraph
        elif in_expertise and line.strip() and not line.startswith("#"):
            # Collect all consecutive expertise lines
            expertise_items = []
            while i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
                item = lines[i].rstrip()
                # Strip bullet markers
                item = _re.sub(r"^[-*•]\s+", "", item).strip()
                item = _re.sub(r"\*+", "", item).strip()
                if item:
                    # If already contains |, split and add items individually
                    if "|" in item:
                        expertise_items.extend([x.strip() for x in item.split("|") if x.strip()])
                    else:
                        expertise_items.append(item)
                i += 1
            if expertise_items:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(" | ".join(expertise_items))
                set_font(run, 10, DARK)
            continue  # i already advanced in the while loop

        # Skills rows – "**Label:**\tValue" bold label, plain value, left tab at 1440
        elif in_skills and not line.startswith("#") and line.strip():
            clean = _re.sub(r"\*+", "", line).strip().rstrip(":")
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)

            # Tab stops matching template: clear 720, left 1620
            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            for val, pos in [("clear", "720"), ("left", "1620")]:
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), val)
                tab.set(qn("w:pos"), pos)
                tabs.append(tab)
            pPr.append(tabs)

            if "\t" in line:
                raw_label = line.split("\t")[0]
                value = line.split("\t", 1)[1]
                label = _re.sub(r"\*+", "", raw_label).strip().rstrip(":")
            elif ":" in clean:
                parts = clean.split(":", 1)
                label, value = parts[0].strip(), parts[1].strip()
            else:
                label, value = clean, ""

            # No colon – tab provides the visual separation
            run1 = p.add_run(label)
            set_font(run1, 10, DARK, bold=True)
            if value:
                run2 = p.add_run("\t" + value.strip())
                set_font(run2, 10, DARK)

        # Education – tab-separated "Degree[TAB]Institution" matching template exactly
        elif in_education and line.strip() and not line.startswith("#"):
            clean = _re.sub(r"^[-*•]\s+", "", line).strip()
            clean = _re.sub(r"\*+", "", clean).strip()
            if not clean:
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)

            # Add tab stops matching template: clear 720, left 3150, right 9026
            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            for val, pos in [("clear", "720"), ("left", "3150"), ("right", "9026")]:
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), val)
                tab.set(qn("w:pos"), pos)
                if val != "clear":
                    tab.set(qn("w:leader"), "none")
                tabs.append(tab)
            pPr.append(tabs)

            # Split on · separator or tab, render as bold degree + tab + plain institution
            if "\t" in clean:
                degree, institution = clean.split("\t", 1)
            elif "  ·  " in clean:
                degree, institution = clean.split("  ·  ", 1)
            elif " · " in clean:
                degree, institution = clean.split(" · ", 1)
            else:
                degree, institution = clean, ""

            run1 = p.add_run(degree.strip())
            set_font(run1, 10, DARK, bold=True)
            if institution:
                run2 = p.add_run("\t" + institution.strip())
                set_font(run2, 10, DARK)

        # Horizontal rule / section break
        elif line.strip() in ("---", "***", "___"):
            pass  # Skip – spacing handled by SB/SA

        # Empty line – skip, spacing handled by SB/SA
        elif line.strip() == "":
            pass

        # Bold inline (areas of expertise etc)
        elif "**" in line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            segments = _re.split(r"(\*\*.*?\*\*)", line)
            for seg in segments:
                if seg.startswith("**") and seg.endswith("**"):
                    run = p.add_run(seg[2:-2])
                    set_font(run, 10, DARK, bold=True)
                else:
                    run = p.add_run(seg)
                    set_font(run, 10, DARK)

        # Plain text (summary, expertise etc)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line.strip())
            set_font(run, 10, DARK)

        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
